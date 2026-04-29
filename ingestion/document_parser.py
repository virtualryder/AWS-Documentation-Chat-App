"""
Extracts plain text from customer-uploaded documents.
Supports: PDF, Word (.docx), plain text (.txt, .md)
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract readable text from an uploaded file.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used to detect format).

    Returns:
        Extracted text string, or an error message if parsing fails.
    """
    ext = Path(filename).suffix.lower()

    try:
        if ext in (".txt", ".md"):
            return file_bytes.decode("utf-8", errors="replace")

        elif ext == ".docx":
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)

        elif ext == ".pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n\n".join(pages)

        else:
            return f"[Unsupported file type: {ext}. Supported: .txt, .md, .docx, .pdf]"

    except Exception as e:
        logger.error("Failed to parse '%s': %s", filename, e)
        return f"[Failed to parse {filename}: {e}]"


def format_uploaded_docs(extracted: list[tuple[str, str]]) -> str:
    """
    Format multiple extracted documents into a single context string.

    Args:
        extracted: List of (filename, text) tuples.

    Returns:
        Formatted string suitable for inclusion in customer context.
    """
    if not extracted:
        return ""
    parts = []
    for filename, text in extracted:
        parts.append(f"### Uploaded Document: {filename}\n{text}")
    return "\n\n---\n\n".join(parts)
